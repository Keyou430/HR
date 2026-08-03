from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "rbac-phase-prompts.md"
OUTPUT = ROOT / "docs" / "rbac-phase-prompts.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK_BLUE = "0B2545"
TABLE_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
GRAY = "666666"
BODY_FONT = "Calibri"
EAST_ASIA_FONT = "Microsoft YaHei"
MONO_FONT = "Consolas"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, name: str, size: float, color: str | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_border(cell, color="B8C3D1", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, BODY_FONT, 9, GRAY)


def set_style_font(style, name=BODY_FONT, size=11, color=None, bold=None):
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def configure_styles(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, size=11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True

    title = doc.styles.add_style("Document Title", 1)
    set_style_font(title, size=24, color=INK_BLUE, bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.line_spacing = 1.0

    subtitle = doc.styles.add_style("Document Subtitle", 1)
    set_style_font(subtitle, size=12, color=GRAY)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(18)

    metadata = doc.styles.add_style("Metadata", 1)
    set_style_font(metadata, size=9, color=GRAY)
    metadata.paragraph_format.space_before = Pt(0)
    metadata.paragraph_format.space_after = Pt(10)

    code = doc.styles.add_style("Code Block", 1)
    set_style_font(code, name=MONO_FONT, size=8.5, color="263238")
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.left_indent = Inches(0.12)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.keep_together = True

    note = doc.styles.add_style("Note Box", 1)
    set_style_font(note, size=10, color=INK_BLUE)
    note.paragraph_format.space_before = Pt(5)
    note.paragraph_format.space_after = Pt(8)
    note.paragraph_format.left_indent = Inches(0.12)
    note.paragraph_format.right_indent = Inches(0.12)

    table_text = doc.styles.add_style("Table Text", 1)
    set_style_font(table_text, size=9.2)
    table_text.paragraph_format.space_before = Pt(0)
    table_text.paragraph_format.space_after = Pt(0)
    table_text.paragraph_format.line_spacing = 1.05

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        set_style_font(style, size=11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("Replica 平台 · RBAC v2.0")
    set_run_font(run, BODY_FONT, 9, GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run("阶段性提示词手册  ·  第 ")
    set_run_font(run, BODY_FONT, 9, GRAY)
    add_page_field(fp)
    run = fp.add_run(" 页")
    set_run_font(run, BODY_FONT, 9, GRAY)


def set_paragraph_shading(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_inline_runs(paragraph, text: str, font_size=11, color=None):
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, BODY_FONT, font_size, color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, BODY_FONT, font_size, color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, MONO_FONT, max(font_size - 0.5, 8.5), "263238")
        else:
            label = token[1 : token.index("](")]
            run = paragraph.add_run(label)
            set_run_font(run, BODY_FONT, font_size, color)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, BODY_FONT, font_size, color)


def add_body_paragraph(doc, text: str, style="Normal"):
    paragraph = doc.add_paragraph(style=style)
    add_inline_runs(paragraph, text)
    return paragraph


def add_code_block(doc, lines: list[str]):
    paragraph = doc.add_paragraph(style="Code Block")
    set_paragraph_shading(paragraph, "F2F4F7")
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(run, MONO_FONT, 8.5, "263238")
        if index != len(lines) - 1:
            run.add_break()
    return paragraph


def parse_table_row(line: str) -> list[str]:
    value = line.strip()
    if value.startswith("|"):
        value = value[1:]
    if value.endswith("|"):
        value = value[:-1]
    return [cell.strip() for cell in value.split("|")]


def is_table_separator(line: str) -> bool:
    cells = parse_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def table_widths(count: int) -> list[int]:
    options = {
        2: [2700, 6660],
        3: [2160, 3600, 3600],
        4: [2016, 2592, 2592, 2160],
        5: [1500, 1900, 2000, 1900, 2060],
    }
    if count in options:
        return options[count]
    base = CONTENT_WIDTH_DXA // max(count, 1)
    widths = [base] * count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_table(doc, rows: list[list[str]]):
    if not rows:
        return
    count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=count)
    widths = table_widths(count)
    set_table_geometry(table, widths)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for row_index, row_data in enumerate(rows):
        row = table.rows[row_index]
        for col_index, cell in enumerate(row.cells):
            set_cell_border(cell)
            if row_index == 0:
                set_cell_shading(cell, TABLE_FILL)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.style = "Table Text"
            if row_index == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(paragraph, row_data[col_index], font_size=9.2, color=INK_BLUE if row_index == 0 else None)
            for run in paragraph.runs:
                if row_index == 0:
                    run.bold = True
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text: str, level: int, first_heading: bool):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    if level == 1 and not first_heading:
        paragraph.paragraph_format.page_break_before = True
    add_inline_runs(paragraph, text, font_size={1: 16, 2: 13, 3: 12}[level], color=BLUE if level < 3 else DARK_BLUE)
    return paragraph


def add_title_block(doc, title: str, metadata_lines: list[str]):
    paragraph = doc.add_paragraph(style="Document Title")
    add_inline_runs(paragraph, title, font_size=24, color=INK_BLUE)
    subtitle = doc.add_paragraph(style="Document Subtitle")
    add_inline_runs(subtitle, "阶段性实施、代码审查与验收提示词手册", font_size=12, color=GRAY)
    for line in metadata_lines:
        paragraph = doc.add_paragraph(style="Metadata")
        add_inline_runs(paragraph, line, font_size=9, color=GRAY)
    note = doc.add_paragraph(style="Note Box")
    add_inline_runs(
        note,
        "使用方式：每次只执行一个 Phase，依次运行实施提示词、阶段审查提示词和阶段验收提示词。",
        font_size=10,
        color=INK_BLUE,
    )
    set_paragraph_shading(note, LIGHT_FILL)


def build_document():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)

    index = 0
    first_h1 = True
    title_consumed = False
    metadata_lines: list[str] = []
    while index < len(lines) and (not lines[index].strip() or lines[index].startswith(">")):
        if lines[index].startswith(">"):
            metadata_lines.append(lines[index][1:].strip())
        index += 1
    if index < len(lines) and lines[index].startswith("# "):
        title = lines[index][2:].strip()
        index += 1
        while index < len(lines) and (not lines[index].strip() or lines[index].startswith(">")):
            if lines[index].startswith(">"):
                metadata_lines.append(lines[index][1:].strip())
            index += 1
        add_title_block(doc, title, metadata_lines)
        title_consumed = True

    in_code = False
    code_lines: list[str] = []
    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if not line.strip():
            index += 1
            continue
        if line.strip() == "---":
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(4)
            index += 1
            continue
        if line.startswith("# "):
            if title_consumed:
                add_heading(doc, line[2:].strip(), 1, first_h1)
                first_h1 = False
            else:
                add_title_block(doc, line[2:].strip(), [])
                title_consumed = True
            index += 1
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), 1, first_h1)
            first_h1 = False
            index += 1
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), 2, True)
            index += 1
            continue
        if line.startswith("#### "):
            add_heading(doc, line[5:].strip(), 3, True)
            index += 1
            continue
        if line.startswith("> "):
            paragraph = doc.add_paragraph(style="Note Box")
            add_inline_runs(paragraph, line[2:].strip(), font_size=10, color=INK_BLUE)
            set_paragraph_shading(paragraph, LIGHT_FILL)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            rows = [parse_table_row(line)]
            index += 2
            while index < len(lines) and lines[index].startswith("|"):
                rows.append(parse_table_row(lines[index]))
                index += 1
            add_table(doc, rows)
            continue
        checkbox = re.match(r"^\s*-\s+\[([ xX])\]\s+(.*)$", line)
        if checkbox:
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, f"☐ {checkbox.group(2)}")
            index += 1
            continue
        bullet = re.match(r"^\s*[-*]\s+(.*)$", line)
        if bullet:
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, bullet.group(1))
            index += 1
            continue
        numbered = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if numbered:
            paragraph = doc.add_paragraph(style="List Number")
            add_inline_runs(paragraph, numbered.group(1))
            index += 1
            continue

        add_body_paragraph(doc, line)
        index += 1

    doc.core_properties.title = "Replica RBAC v2.0 分阶段实施、审查与验收提示词"
    doc.core_properties.subject = "Replica 平台权限系统实施工作文档"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "Replica, RBAC, ABAC, AI Security, Implementation Prompts"
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
